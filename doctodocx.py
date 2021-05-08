import win32com.client
import os
import docx

word = win32com.client.Dispatch("Word.Application")
word.visible = False

wordfile = word.Documents.Open("C:\\Users\\110501\\Desktop\\proposal\\LCEWA01_100311_00088 (2).doc")
wordfile.SaveAs2("C:\\Users\\110501\\Desktop\\proposal\\LCEWA01_100311_00088 (2).doc"+'x', FileFormat = 16) 
#另存新檔,附檔名多添加x,儲存成docx格式代碼
wordfile.Close()

doc = docx.Document('C:\\Users\\110501\\Desktop\\proposal\\LCEWA01_100311_00088 (2).docx')
print('run數量：', len(doc.paragraphs))

for para in doc.paragraphs:
    print(para.text)
'''for para in doc.paragraphs:
    if para.style.name.startswith('總說明'):
        print(para.text)'''